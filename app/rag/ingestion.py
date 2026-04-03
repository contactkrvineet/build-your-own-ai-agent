"""
Document ingestion — loads PDF, TXT, DOCX, MD files from the configured
documents folder and returns LangChain Document objects.

Supports:
  - Native PDF text extraction (pypdf)
  - OCR fallback for scanned PDFs (pytesseract + pdf2image)
  - DOCX (python-docx via LangChain loader)
  - Markdown and plain text
"""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import List, Optional, Set

from langchain_core.documents import Document

from app.config.settings import get_settings
from app.utils.logger import logger


# ---------------------------------------------------------------------------
# Supported file types
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS: Set[str] = {".pdf", ".txt", ".docx", ".md"}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def load_documents(
    folder_path: Optional[str] = None,
    ocr_enabled: Optional[bool] = None,
    ocr_language: Optional[str] = None,
) -> List[Document]:
    """
    Walk *folder_path* and load every supported document.

    Args:
        folder_path:  Override config.rag.documents_path.
        ocr_enabled:  Override config.rag.ocr_enabled.
        ocr_language: Override config.rag.ocr_language.

    Returns:
        List of LangChain Document objects with populated metadata.
    """
    s = get_settings()
    path = Path(folder_path or s.rag_documents_path)
    use_ocr = ocr_enabled if ocr_enabled is not None else s.rag_ocr_enabled
    lang = ocr_language or s.rag_ocr_language

    if not path.exists():
        logger.warning(f"Documents folder not found: {path}. Creating it.")
        path.mkdir(parents=True, exist_ok=True)
        return []

    docs: List[Document] = []
    files = [
        f for f in path.iterdir()
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    ]

    if not files:
        logger.info(f"No supported documents found in {path}")
        return []

    logger.info(f"Ingesting {len(files)} documents from {path}")

    for file_path in files:
        try:
            loaded = _load_single_file(file_path, use_ocr, lang)
            docs.extend(loaded)
            logger.debug(f"Loaded {len(loaded)} page(s) from {file_path.name}")
        except Exception as e:
            logger.error(f"Failed to load {file_path.name}: {e}")

    logger.info(f"Total pages/chunks ingested: {len(docs)}")
    return docs


def load_single_document(
    file_path: str | Path,
    ocr_enabled: Optional[bool] = None,
    ocr_language: Optional[str] = None,
) -> List[Document]:
    """Load a single file — used by the hot-reload file watcher."""
    s = get_settings()
    use_ocr = ocr_enabled if ocr_enabled is not None else s.rag_ocr_enabled
    lang = ocr_language or s.rag_ocr_language
    return _load_single_file(Path(file_path), use_ocr, lang)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _load_single_file(
    file_path: Path, ocr_enabled: bool, ocr_language: str
) -> List[Document]:
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(file_path, ocr_enabled, ocr_language)
    elif suffix == ".txt":
        return _load_text(file_path)
    elif suffix == ".docx":
        return _load_docx(file_path)
    elif suffix == ".md":
        return _load_markdown(file_path)
    else:
        logger.warning(f"Skipping unsupported file: {file_path.name}")
        return []


def _file_hash(file_path: Path) -> str:
    """MD5 hash of file content for deduplication."""
    h = hashlib.md5()
    with file_path.open("rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def _base_metadata(file_path: Path) -> dict:
    return {
        "source": str(file_path),
        "filename": file_path.name,
        "extension": file_path.suffix.lower(),
        "file_hash": _file_hash(file_path),
    }


def _load_pdf(file_path: Path, ocr_enabled: bool, ocr_language: str) -> List[Document]:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    docs: List[Document] = []
    meta = _base_metadata(file_path)

    for page_num, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()

        # Fallback to OCR if page has no extractable text
        if not text and ocr_enabled:
            logger.debug(f"Page {page_num+1} of {file_path.name} is scanned — using OCR")
            from app.utils.ocr import extract_text_with_ocr

            text = _ocr_single_page(file_path, page_num, ocr_language) or ""

        if text:
            docs.append(
                Document(
                    page_content=text,
                    metadata={**meta, "page": page_num + 1},
                )
            )

    return docs


def _ocr_single_page(file_path: Path, page_num: int, language: str) -> Optional[str]:
    """OCR a single PDF page by converting it to an image first."""
    try:
        from pdf2image import convert_from_path
        import pytesseract

        images = convert_from_path(
            str(file_path), dpi=300, first_page=page_num + 1, last_page=page_num + 1
        )
        if images:
            return pytesseract.image_to_string(images[0], lang=language)
    except Exception as e:
        logger.warning(f"OCR failed for page {page_num+1} of {file_path.name}: {e}")
    return None


def _load_text(file_path: Path) -> List[Document]:
    text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        return []
    return [Document(page_content=text, metadata=_base_metadata(file_path))]


def _load_docx(file_path: Path) -> List[Document]:
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        if not text:
            return []
        return [Document(page_content=text, metadata=_base_metadata(file_path))]
    except Exception as e:
        logger.error(f"DOCX load error for {file_path.name}: {e}")
        return []


def _load_markdown(file_path: Path) -> List[Document]:
    text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        return []
    return [Document(page_content=text, metadata=_base_metadata(file_path))]
